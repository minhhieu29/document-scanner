import customtkinter as ctk
from tkinter import filedialog, messagebox, ttk, Menu
import threading
import os
import hashlib
from datetime import datetime
from pathlib import Path
import docx  
import fitz  
from utils import export_to_excel
from concurrent.futures import ThreadPoolExecutor 
import google.generativeai as genai
import win32com.client

class DocumentScanner:
    def __init__(self):
        self.extensions = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.pptx', '.txt', '.md'}
        self.exclude_dirs = {'Windows', 'AppData', 'Program Files', 'node_modules', '$Recycle.Bin', 'System Volume Information'}

    def get_file_hash(self, file_path):
        hash_md5 = hashlib.md5()
        try:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except: return None

    def read_file_content(self, file_path):
        content = ""
        ext = Path(file_path).suffix.lower()
        try:
            # 1. Đọc file Text và Markdown
            if ext in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            
            # 2. Đọc file Word mới (.docx)
            elif ext == '.docx':
                doc = docx.Document(file_path)
                full_text = [para.text for para in doc.paragraphs]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text.append(cell.text)
                content = "\n".join(full_text)
            
            # 3. Đọc file Word cũ (.doc) - Cần cài pywin32
            elif ext == '.doc':
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False # Chạy ngầm không hiện cửa sổ Word
                doc = word.Documents.Open(str(Path(file_path).absolute()))
                content = doc.Content.Text
                doc.Close()
                word.Quit()
            
            # 4. Đọc file PDF
            elif ext == '.pdf':
                with fitz.open(file_path) as doc:
                    content = "".join([page.get_text() for page in doc])
                    
        except Exception as e: 
            print(f"Lỗi khi đọc file {file_path}: {e}")
            pass
            
        return content.lower()

    def scan_directory(self, root_path):
        all_files = []
        path_obj = Path(root_path)
        for entry in path_obj.rglob('*'):
            try:
                if entry.is_file() and entry.suffix.lower() in self.extensions:
                    if not any(ex in entry.parts for ex in self.exclude_dirs):
                        all_files.append(entry)
            except: continue
            
        def process_file(entry):
            return {
                "name": entry.name,
                "path": str(entry.absolute()),
                "size_display": f"{entry.stat().st_size / 1024:.2f} KB",
                "modified": datetime.fromtimestamp(entry.stat().st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
                "hash": self.get_file_hash(entry)
            }

        with ThreadPoolExecutor(max_workers=8) as executor:
            results = list(executor.map(process_file, all_files))
        return results

class App(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("Document Scanner Pro - Final Edition")
        self.geometry("1100x700")
        ctk.set_appearance_mode("Dark")
        
        self.scanner = DocumentScanner()
        self.results = [] 

        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # Sidebar
        self.sidebar = ctk.CTkFrame(self, width=220, corner_radius=0)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        
        self.logo_label = ctk.CTkLabel(self.sidebar, text="SCANNER", font=ctk.CTkFont(size=24, weight="bold"))
        self.logo_label.grid(row=0, column=0, padx=20, pady=(30, 20))

        self.select_btn = ctk.CTkButton(self.sidebar, text="Chọn thư mục", command=self.select_path, fg_color="#6272a4")
        self.select_btn.grid(row=1, column=0, padx=20, pady=10)

        self.scan_btn = ctk.CTkButton(self.sidebar, text="Bắt đầu Quét", fg_color="#ff79c6", hover_color="#bd93f9", command=self.start_scan_thread)
        self.scan_btn.grid(row=2, column=0, padx=20, pady=10)

        self.export_btn = ctk.CTkButton(self.sidebar, text="Xuất Excel", state="disabled", command=self.export_data)
        self.export_btn.grid(row=3, column=0, padx=20, pady=10)

        # Main Content
        self.main_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.main_frame.grid(row=0, column=1, padx=20, pady=20, sticky="nsew")
        self.main_frame.grid_columnconfigure(0, weight=1)
        self.main_frame.grid_rowconfigure(3, weight=1)

        self.search_entry = ctk.CTkEntry(self.main_frame, placeholder_text="Nhập từ khóa...", height=35, border_color="#ff79c6")
        self.search_entry.grid(row=0, column=0, padx=10, pady=(0, 5), sticky="ew")
        self.search_entry.bind("<Return>", self.filter_results)

        self.content_search_var = ctk.BooleanVar(value=False)
        self.content_check = ctk.CTkCheckBox(self.main_frame, text="Tìm trong nội dung file (Deep Scan)", variable=self.content_search_var)
        self.content_check.grid(row=1, column=0, padx=10, pady=(0, 10), sticky="w")

        self.status_label = ctk.CTkLabel(self.main_frame, text="Sẵn sàng quét.", text_color="#ff79c6")
        self.status_label.grid(row=2, column=0, padx=10, pady=5, sticky="w")

        self.setup_treeview()

    def setup_treeview(self):
        style = ttk.Style()
        style.theme_use("default")
        style.configure("Treeview", 
                        background="#252526", 
                        foreground="#cccccc", 
                        fieldbackground="#252526", 
                        rowheight=35, 
                        font=("Segoe UI", 12)) 
        
        style.configure("Treeview.Heading", 
                        background="#2d2d2d", 
                        foreground="white", 
                        relief="flat", 
                        font=("Segoe UI", 13, "bold"))
       

        style.map("Treeview", background=[('selected', "#007acc")])

        tree_frame = ctk.CTkFrame(self.main_frame, corner_radius=10)
        tree_frame.grid(row=3, column=0, sticky="nsew")
        tree_frame.grid_columnconfigure(0, weight=1)
        tree_frame.grid_rowconfigure(0, weight=1)

        columns = ("no", "name", "size", "modified")
        self.tree = ttk.Treeview(tree_frame, columns=columns, show='headings', style="Treeview")
        
        self.tree.heading("no", text="STT")
        self.tree.heading("name", text=" Tên File")
        self.tree.heading("size", text="Dung lượng")
        self.tree.heading("modified", text="Ngày sửa")

        self.tree.column("no", width=60, anchor="center")
        self.tree.column("name", width=450, anchor="w")
        self.tree.column("size", width=120, anchor="center")
        self.tree.column("modified", width=180, anchor="center")

        self.tree.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        self.tree.bind("<Double-1>", self.open_file_event)
        self.tree.bind("<Button-3>", self.show_context_menu)

        self.context_menu = Menu(self, tearoff=0, bg="#2d2d2d", fg="white")
        self.context_menu.add_command(label="Mở File", command=self.open_file_event)
        self.context_menu.add_command(label="Mở thư mục chứa file", command=self.open_folder_location)
        self.context_menu.add_command(label="Copy đường dẫn", command=self.copy_file_path)
        self.context_menu.add_separator() 
        self.context_menu.add_command(label=" Tóm tắt nội dung (AI)", command=self.ai_summarize_event)

        self.scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscroll=self.scrollbar.set)
        self.scrollbar.grid(row=0, column=1, sticky="ns", pady=10)

    # HÀM XỬ LÝ 
    def show_context_menu(self, event):
        item = self.tree.identify_row(event.y)
        if item:
            self.tree.selection_set(item)
            self.context_menu.post(event.x_root, event.y_root)

    def open_file_event(self, event=None):
        selected = self.tree.selection()
        if not selected: return
        item_data = self.tree.item(selected)
        file_name = item_data['values'][1].strip()
        for res in self.results:
            if res['name'] == file_name:
                os.startfile(res['path'])
                break

    def open_folder_location(self):
        selected = self.tree.selection()
        if not selected: return
        item_data = self.tree.item(selected)
        file_name = item_data['values'][1].strip()
        for res in self.results:
            if res['name'] == file_name:
                os.startfile(os.path.dirname(res['path']))
                break

    def copy_file_path(self):
        selected = self.tree.selection()
        if not selected: return
        item_data = self.tree.item(selected)
        file_name = item_data['values'][1].strip()
        for res in self.results:
            if res['name'] == file_name:
                self.clipboard_clear()
                self.clipboard_append(res['path'])
                messagebox.showinfo("OK", "Đã copy đường dẫn!")
                break

    def select_path(self):
        path = filedialog.askdirectory()
        if path:
            self.selected_path = path
            self.status_label.configure(text=f"Đã chọn: {path}")

    def start_scan_thread(self):
        if hasattr(self, 'selected_path'):
            self.scan_btn.configure(state="disabled")
            for item in self.tree.get_children(): self.tree.delete(item)
            threading.Thread(target=self.run_scan, daemon=True).start()
        else:
            messagebox.showwarning("!", "Chưa chọn thư mục!")

    def run_scan(self):
        self.status_label.configure(text="Đang quét đa luồng...")
        self.results = self.scanner.scan_directory(self.selected_path)
        self.update_table(self.results)
        self.status_label.configure(text=f"Xong! Tìm thấy {len(self.results)} file.")
        self.scan_btn.configure(state="normal")
        self.export_btn.configure(state="normal")

    def update_table(self, data_list):
        for item in self.tree.get_children(): self.tree.delete(item)
        for i, item in enumerate(data_list, start=1):
            self.tree.insert("", "end", values=(i, f" {item['name']}", item['size_display'], item['modified']))

    def filter_results(self, event=None):
        query = self.search_entry.get().lower()
        if not query:
            self.update_table(self.results)
            return
        self.status_label.configure(text="Đang lục soát...")
        self.update_idletasks()
        filtered = []
        search_content = self.content_search_var.get()
        for item in self.results:
            if query in item['name'].lower() or (search_content and query in self.scanner.read_file_content(item['path'])):
                filtered.append(item)
        self.update_table(filtered)
        self.status_label.configure(text=f"Tìm thấy {len(filtered)} kết quả.")

    def export_data(self):
        export_to_excel(self.results, "report.xlsx")
        messagebox.showinfo("Xong", "Đã xuất Excel!")

    def ai_summarize_event(self):
        selected = self.tree.selection()
        if not selected: return
        
        
        item_data = self.tree.item(selected)
        file_name = item_data['values'][1].strip()
        file_path = ""
        for res in self.results:
            if res['name'] == file_name:
                file_path = res['path']
                break
        
        if not file_path: return

        
        content = self.scanner.read_file_content(file_path)
        if not content or len(content) < 20:
            messagebox.showwarning("!", "File không có nội dung hoặc quá ngắn để tóm tắt!")
            return

        
        self.status_label.configure(text=" AI đang đọc và tóm tắt...")
        threading.Thread(target=self.run_ai_logic, args=(content,), daemon=True).start()

    def run_ai_logic(self, content):
        try:
            genai.configure(api_key="AIzaSyBbWVgkn9zUUBuJu7MNKgV6JIfFcu3oBbA") 
            
            model = genai.GenerativeModel('gemini-2.5-flash')
            
            prompt = f"Hãy tóm tắt nội dung sau đây một cách ngắn gọn, súc tích bằng tiếng Việt (khoảng 3-5 ý chính):\n\n{content[:8000]}" 
            
            response = model.generate_content(prompt)
            
            if response.text:
                self.show_ai_result(response.text)
                self.status_label.configure(text="✨ AI: Tóm tắt hoàn tất!")
            else:
                messagebox.showwarning("AI", "AI không trả về nội dung, thử lại file khác xem.")
                
        except Exception as e:
            error_msg = str(e)
            messagebox.showerror("Lỗi AI", f"Chi tiết lỗi: {error_msg}")
            self.status_label.configure(text="Lỗi kết nối AI.")

    def show_ai_result(self, text):
        top = ctk.CTkToplevel(self)
        top.title("Kết quả tóm tắt AI")
        top.geometry("600x400")
        top.attributes("-topmost", True)

        txt_area = ctk.CTkTextbox(top, font=("Segoe UI", 13), corner_radius=10)
        txt_area.pack(padx=20, pady=20, fill="both", expand=True)
        txt_area.insert("0.0", text)
        txt_area.configure(state="disabled")

if __name__ == "__main__":
    app = App()
    app.mainloop()